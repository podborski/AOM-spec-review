"""
The Clear BSD License

Copyright (c) 2023, Apple Inc.

All rights reserved.

Redistribution and use in source and binary forms, with or without modification,
are permitted (subject to the limitations in the disclaimer below) provided that
the following conditions are met:

Redistributions of source code must retain the above copyright notice, this list
of conditions and the following disclaimer.

Redistributions in binary form must reproduce the above copyright notice, this
list of conditions and the following disclaimer in the documentation and/or other
materials provided with the distribution.

Neither the name of the Apple Inc. nor the names of its contributors
may be used to endorse or promote products derived from this software without
specific prior written permission.

NO EXPRESS OR IMPLIED LICENSES TO ANY PARTY'S PATENT RIGHTS ARE GRANTED BY THIS
LICENSE. THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE DISCLAIMED.
IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE FOR ANY DIRECT,
INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES (INCLUDING, BUT
NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR
PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY,
WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)
ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE
POSSIBILITY OF SUCH DAMAGE.
"""

import os
import re
import sys
import hashlib
import subprocess
from docx import Document
from loguru import logger
from github import Github, Auth
from argparse import ArgumentParser

from issue_automation.helpers import RateLimitRetry, add_hyperlink

LABELS = {
    "ed": "editorial",
    "ge": "general",
    "te": "technical",
    "?": "question",
}


def assert_log(condition, message):
    if not condition:
        logger.critical(message)
        raise AssertionError(message)


def check_if_github_cli_is_installed():
    # Check if Github CLI is installed
    try:
        subprocess.check_output(["gh", "--version"])
    except FileNotFoundError:
        logger.critical(
            "Github CLI is not installed. Please install it from https://cli.github.com"
        )
        sys.exit(1)


def create_issue_meta(rows):
    # Add header
    header = ["Version", "Source", "Clause(s)"]
    rows.insert(0, header)

    # Create table
    table = ""
    for i, row in enumerate(rows):
        table += "|"
        table += "|".join(row)
        table += "|\n"
        if i == 0:
            table += "|"
            table += "|".join([":---:"] * len(row))
            table += "|\n"
    return table


def get_comments(table):
    # There should be 6 columns
    assert_log(len(table.columns) == 6, "Table should have exactly 6 columns")
    assert_log(len(table.rows) > 1, "Table should have at least one row")

    comments = []
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        cell_values = []
        for cell in row.cells:
            parsed_text = ""
            for paragraph in cell.paragraphs:
                if "spec" in paragraph.style.name:
                    parsed_text += f"<blockquote>\n{paragraph.text}\n</blockquote>\n\n"
                    continue

                if "code" in paragraph.style.name:
                    parsed_text += f"\n```\n{paragraph.text}\n```\n"
                    continue

                for run in paragraph.runs:
                    if "code" in run.style.name:
                        parsed_text += f"`{run.text}`"
                    else:
                        parsed_text += run.text
                parsed_text += "\n"
            cell_values.append(parsed_text.strip())
        if any(c != "" for c in cell_values):
            comments.append(cell_values)
    return comments


def create_issues(comments, repo, version, existing_ids, args):
    created_issues = {}
    issues_created = 0
    issues_skipped = 0
    for i, row in enumerate(comments):
        # Process labels
        labels = []
        if row[0] != "":
            for label in row[0].split(","):
                s_label = label.strip()
                assert_log(s_label in LABELS, f"Invalid label: {s_label}")
                labels.append(LABELS[s_label])

        # Check if type is in type_filter
        if args.type_filter:
            if row[0] == "":
                match = "all" in args.type_filter
            else:
                match = any(
                    t.strip().lower() in row[0].lower()
                    for t in args.type_filter.split(",")
                )

            if not match:
                logger.info(
                    f"Skipping {row[3]} as it is not in type(s) [{args.type_filter}]"
                )
                issues_skipped += 1
                continue

        # Process clauses and title
        clause = None
        if row[2] != "":
            clause = [f"§{c.strip()}" for c in row[2].split(",")]
            clause = ", ".join(clause)
            title = f"{clause}: {row[3]}"
        else:
            title = row[3]

        # Check if clause is in clause_filter
        if args.clause_filter:
            if not clause:
                match = "all" in args.clause_filter
            else:
                match = any(
                    re.search(rf"(?<!\.){c.strip()}", clause)
                    for c in args.clause_filter.split(",")
                )

            if not match:
                logger.info(
                    f"Skipping {row[3]} as it is not in clause(s) [{args.clause_filter}]"
                )
                issues_skipped += 1
                continue

        # If comment is internal, skip
        if clause and "!" in clause:
            logger.warning(f"Skipping internal comment: {row[3]}")
            continue

        # Get rest of the data
        source = row[1]
        comment = row[4]
        suggestion = row[5]

        if source == "":
            source = "Unknown"
            logger.warning(f"No source for: {title}, setting to 'Unknown'")

        raw_id = f"{title}{comment}{suggestion}"
        id_hash = hashlib.sha1(raw_id.encode("utf-8")).hexdigest()

        if id_hash in existing_ids:
            logger.info(f"Skipping existing issue: {title}")
            issues_skipped += 1
            continue

        body = f"<!-- id: {id_hash} -->\n"
        body += create_issue_meta([[version, source, clause or "all"]])
        body += "\n"

        if comment == "":
            logger.warning(f"No comment for: {title}, skipping...")
            issues_skipped += 1
            continue
        body += f"\n#### Comment:\n{comment}\n"

        if suggestion != "":
            body += f"\n-----\n#### Suggestion:\n{suggestion}\n"

        if args.limit and issues_created >= args.limit:
            logger.info(f"Reached limit of {args.limit} issues")
            break

        if args.dry_run:
            logger.info(f"Would create issue: {title}")
            logger.debug(f"\n{body}")
            issues_skipped += 1
        else:
            logger.success(f"Creating issue: {title}")
            issue = repo.create_issue(title=title, body=body, labels=labels)
            issues_created += 1
            created_issues[i] = issue.html_url

    logger.success(f"created={issues_created} skipped={issues_skipped}")
    return created_issues


def process_comments_document():
    # Change logger format
    logger.remove()  # All configured handlers are removed
    fmt = "<green>{time:HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>"
    logger.add(sys.stderr, format=fmt)

    parser = ArgumentParser(description="Process comments document")
    parser.add_argument(
        "-i", help="Path to comments document", dest="comments_document", required=True
    )
    parser.add_argument("-l", "--limit", help="Limit number of issues", type=int)
    parser.add_argument(
        "-t",
        "--type-filter",
        help="Only process type, either string or 'all'. Must be separated by a comma",
        type=str,
    )
    parser.add_argument(
        "-c",
        "--clause-filter",
        help="Only process clause, either number or 'all'. Must be separated by a comma",
        type=str,
    )
    parser.add_argument(
        "--link-titles",
        help="Link comment titles to created issues",
        action="store_true",
    )
    parser.add_argument(
        "--output-document",
        help="Output document with links embedded to title",
        dest="output_document",
    )
    parser.add_argument("-n", "--dry_run", help="Dry run", action="store_true")

    # Check arguments
    args = parser.parse_args()
    assert_log(
        not args.link_titles ^ (args.output_document is not None),
        "Both --link-titles and --output-document must be set",
    )
    assert_log(
        os.path.exists(args.comments_document),
        f"Comments document does not exist: {args.comments_document}",
    )

    # Check if access token is set
    check_if_github_cli_is_installed()
    auth_token = os.popen("gh auth token").read().strip()

    # Initialize Github
    auth = Auth.Token(auth_token)
    git = Github(auth=auth, retry=RateLimitRetry(total=10, backoff_factor=0.1))

    document = Document(args.comments_document)

    # Get GitHub repository and version from header
    version = document.sections[0].header.tables[0].rows[1].cells[1].text
    github_repo = document.sections[0].header.tables[0].rows[0].cells[2].text
    github_repo = re.search(r"github\.com\/(.+)\s*", github_repo).group(1)
    repo = git.get_repo(github_repo)

    # Get existing issues
    all_issues = repo.get_issues(state="all")
    existing_ids = set()
    for issue in all_issues:
        if issue.body is None:
            continue
        issue_id = re.search(r"<!-- id: (.+) -->", issue.body)
        if issue_id:
            issue_id = issue_id.group(1)
            existing_ids.add(issue_id)

    # Process table
    assert_log(len(document.tables) == 1, "Document should have exactly one table")
    table = document.tables[0]
    comments = get_comments(table)

    logger.info(f"Found {len(comments)} comments")

    # Create issues
    issues = create_issues(comments, repo, version, existing_ids, args)

    # Add links to titles
    if not args.link_titles or args.dry_run or len(issues.keys()) == 0:
        return

    logger.info("Adding links to titles")

    # If output document already exists, update it
    if os.path.exists(args.output_document):
        document = Document(args.output_document)
        table = document.tables[0]

    for i, row in enumerate(table.rows[1:]):
        if i not in issues:
            continue
        cell = row.cells[3]
        add_hyperlink(cell.paragraphs[0], issues[i], cell.text)

    # Save document
    document.save(args.output_document)
    logger.success(f"Saved document to {args.output_document}")
